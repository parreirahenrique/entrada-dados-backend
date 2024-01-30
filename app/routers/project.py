from fastapi import status, HTTPException, Depends, APIRouter
from sqlalchemy.orm import Session
from typing import List, Optional
from .. import models, schemas, oauth2, utils
from ..database import get_db
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from datetime import date
import docx, os, warnings
from docx.shared import Inches


router = APIRouter(
    tags=['Projects']
)

# Função para inserir um novo módulo
@router.post('/projects', status_code=status.HTTP_201_CREATED, response_model=schemas.ProjectOut)
def create_project(projeto_inserido: schemas.ProjectInsert, db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        projetos = db.query(models.Projeto).all()

        # Determinando a maior ID dentro dos projetos
        id_max: int = 0
        for projeto in projetos:
            if projeto.id > id_max:
                id_max = projeto.id

        # Adicionando novo projeto ao banco de dados
        novo_projeto = models.Projeto(**projeto_inserido.dict())
        novo_projeto.criado_por = usuario_atual.id
        novo_projeto.id = id_max + 1

        # Determinando carga da instalação
        novo_projeto.carga = utils.carga_instalacao(novo_projeto.aumento_carga, novo_projeto.n_fases, novo_projeto.disjuntor, novo_projeto.novo_n_fases, novo_projeto.novo_disjuntor)
        
        # Determinando a potência do módulo 1
        modulo_1 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto_inserido.modelo_modulo_1).first()
        novo_projeto.potencia_modulo_1 = modulo_1.potencia

        # Determinando a potência do módulo 2
        novo_projeto.modulo_anterior_2 = False
        novo_projeto.quantidade_modulo_2 = 0
        novo_projeto.potencia_modulo_2 = 0
        
        if projeto_inserido.modelo_modulo_2 != None:
            modulo_2 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto_inserido.modelo_modulo_2).first()
            novo_projeto.modulo_anterior_2 = projeto_inserido.modulo_anterior_2
            novo_projeto.quantidade_modulo_2 = projeto_inserido.quantidade_modulo_2
            novo_projeto.potencia_modulo_2 = modulo_2.potencia
            
        # Determinando a quantidade total de módulos
        novo_projeto.quantidade_modulos = novo_projeto.quantidade_modulo_1 + novo_projeto.quantidade_modulo_2
        
        # Determinando a potência total dos módulos anteriores
        novo_projeto.potencia_modulos_anterior = utils.potencia_total_modulos_anterior(
            novo_projeto.modulo_anterior_1, novo_projeto.quantidade_modulo_1, novo_projeto.potencia_modulo_1,
            novo_projeto.modulo_anterior_2, novo_projeto.quantidade_modulo_2, novo_projeto.potencia_modulo_2
        )
        
        # Determinando a potência total dos módulos
        novo_projeto.potencia_modulos = utils.potencia_total_modulos(
            novo_projeto.quantidade_modulo_1, novo_projeto.potencia_modulo_1,
            novo_projeto.quantidade_modulo_2, novo_projeto.potencia_modulo_2
        )
        
        # Determinando a área total ocupada pelos módulos
        if projeto_inserido.modelo_modulo_2 != None:
            novo_projeto.area = utils.area_modulos(
                novo_projeto.quantidade_modulo_1, modulo_1.comprimento, modulo_1.largura,
                novo_projeto.quantidade_modulo_2, modulo_2.comprimento, modulo_2.largura
            )

        else:
            novo_projeto.area = utils.area_modulos(
                novo_projeto.quantidade_modulo_1, modulo_1.comprimento, modulo_1.largura,
                0, 0, 0
            )
        
        # Determinando a potência do inversor 1
        inversor_1 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto_inserido.modelo_inversor_1).first()
        novo_projeto.potencia_inversor_1 = inversor_1.potencia

        # Determinando a potência do inversor 2
        novo_projeto.inversor_anterior_2 = False
        novo_projeto.quantidade_inversor_2 = 0
        novo_projeto.potencia_inversor_2 = 0

        if projeto_inserido.modelo_inversor_2 != None:
            inversor_2 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto_inserido.modelo_inversor_2).first()
            novo_projeto.inversor_anterior_2 = projeto_inserido.inversor_anterior_2
            novo_projeto.quantidade_inversor_2 = projeto_inserido.quantidade_inversor_2
            novo_projeto.potencia_inversor_2 = inversor_2.potencia

        # Determinando a potência do inversor 3
        novo_projeto.inversor_anterior_3 = False
        novo_projeto.quantidade_inversor_3 = 0
        novo_projeto.potencia_inversor_3 = 0

        if projeto_inserido.modelo_inversor_3 != None:
            inversor_3 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto_inserido.modelo_inversor_3).first()
            novo_projeto.inversor_anterior_3 = projeto_inserido.inversor_anterior_3
            novo_projeto.quantidade_inversor_3 = projeto_inserido.quantidade_inversor_3
            novo_projeto.potencia_inversor_3 = inversor_3.potencia

        # Determinando a potência do inversor 4
        novo_projeto.inversor_anterior_4 = False
        novo_projeto.quantidade_inversor_4 = 0
        novo_projeto.potencia_inversor_4 = 0

        if projeto_inserido.modelo_inversor_4 != None:
            inversor_4 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto_inserido.modelo_inversor_4).first()
            novo_projeto.inversor_anterior_4 = projeto_inserido.inversor_anterior_4
            novo_projeto.quantidade_inversor_4 = projeto_inserido.quantidade_inversor_4
            novo_projeto.potencia_inversor_4 = inversor_4.potencia

        # Determinando a quantidade total de inversores
        novo_projeto.quantidade_inversores = projeto_inserido.quantidade_inversor_1 + projeto_inserido.quantidade_inversor_2 + projeto_inserido.quantidade_inversor_3 + projeto_inserido.quantidade_inversor_4
        
        # Determinando a potência total dos módulos anteriores
        novo_projeto.potencia_inversores_anterior = utils.potencia_total_inversores_anterior(
            novo_projeto.inversor_anterior_1, novo_projeto.quantidade_inversor_1, novo_projeto.potencia_inversor_1,
            novo_projeto.inversor_anterior_2, novo_projeto.quantidade_inversor_2, novo_projeto.potencia_inversor_2,
            novo_projeto.inversor_anterior_3, novo_projeto.quantidade_inversor_3, novo_projeto.potencia_inversor_3,
            novo_projeto.inversor_anterior_4, novo_projeto.quantidade_inversor_4, novo_projeto.potencia_inversor_4
        )

        # Determinando a potência total dos módulos anteriores
        novo_projeto.potencia_inversores = utils.potencia_total_inversores(
            novo_projeto.quantidade_inversor_1, novo_projeto.potencia_inversor_1,
            novo_projeto.quantidade_inversor_2, novo_projeto.potencia_inversor_2,
            novo_projeto.quantidade_inversor_3, novo_projeto.potencia_inversor_3,
            novo_projeto.quantidade_inversor_4, novo_projeto.potencia_inversor_4
        )
        
        # Determinando a geração mensal
        cidade = db.query(models.Instalacao).filter(models.Instalacao.numero_instalacao == novo_projeto.numero_instalacao).first().cidade
        novo_projeto.geracao_1 = utils.geracao(novo_projeto.potencia_modulos, cidade, 1)
        novo_projeto.geracao_2 = utils.geracao(novo_projeto.potencia_modulos, cidade, 2)
        novo_projeto.geracao_3 = utils.geracao(novo_projeto.potencia_modulos, cidade, 3)
        novo_projeto.geracao_4 = utils.geracao(novo_projeto.potencia_modulos, cidade, 4)
        novo_projeto.geracao_5 = utils.geracao(novo_projeto.potencia_modulos, cidade, 5)
        novo_projeto.geracao_6 = utils.geracao(novo_projeto.potencia_modulos, cidade, 6)
        novo_projeto.geracao_7 = utils.geracao(novo_projeto.potencia_modulos, cidade, 7)
        novo_projeto.geracao_8 = utils.geracao(novo_projeto.potencia_modulos, cidade, 8)
        novo_projeto.geracao_9 = utils.geracao(novo_projeto.potencia_modulos, cidade, 9)
        novo_projeto.geracao_10 = utils.geracao(novo_projeto.potencia_modulos, cidade, 10)
        novo_projeto.geracao_11 = utils.geracao(novo_projeto.potencia_modulos, cidade, 11)
        novo_projeto.geracao_12 = utils.geracao(novo_projeto.potencia_modulos, cidade, 12)
        
        db.add(novo_projeto)
        db.commit()
        db.refresh(novo_projeto)
        return novo_projeto
    
    # Caso o usuário atual não esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de inserir um novo módulo')
    
# Função para adquirir os dados de um projeto
@router.get('/projects/{id}', response_model=schemas.ProjectOut)
def get_one_project(id: int, db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        # Filtrando o projeto do banco de dados de acordo com a ID informada
        projeto = db.query(models.Projeto).filter(models.Projeto.id == id).first()
        
        # Caso a ID informada esteja presente no banco de dados
        if projeto != None:
            return projeto

        # Caso a ID informada não esteja presente no banco de dados
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Projeto com a ID {id} não foi encontrado')
        
    # Caso o usuário atual não esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de adquirir as informações de um projeto')
    
# Função para adquirir os dados de todos projetos
@router.get('/projects', response_model=List[schemas.ProjectOut])
def get_projects(db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual), limite: int = 10, pular: int = 0, buscar: Optional[int] = 0):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        # Filtrando os projetos do banco de dados de acordo com os parâmetros passados, caso tenha sido buscado algum número de instalação
        if buscar != 0:
            projects = db.query(models.Projeto).filter(models.Projeto.numero_instalacao == buscar).limit(limite).offset(pular).all()

        # Filtrando os projetos do banco de dados de acordo com os parâmetros passados, caso não tenha sido buscado algum número de instalação
        elif buscar == 0:
            projects = db.query(models.Projeto).limit(limite).offset(pular).all()
        
        # Caso os parâmetros de busca tenham sido válidos
        if projects != []:
            return projects
        
        # Caso os parâmetros de busca não tenham sido válidos
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Parâmetros de busca passados não exibiram nenhum resultado')
    
    # Caso o usuário atual não esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de adquirir os dados dos módulos')
    
# Função para adquirir os dados de todos projetos
@router.get('/all-projects', response_model=List[schemas.ProjectOut])
def get_projects(db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        projects = db.query(models.Projeto).all()
        
        # Caso os parâmetros de busca tenham sido válidos
        if projects != []:
            return projects
        
        # Caso os parâmetros de busca não tenham sido válidos
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Parâmetros de busca passados não exibiram nenhum resultado')
    
    # Caso o usuário atual não esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de adquirir os dados dos módulos')
    
# Função para atualizar os campos de um projeto
@router.patch('/projects/{id}', response_model=schemas.ProjectOut)
def update_project(id: int, projeto_atualizado: schemas.ProjectUpdate, db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        projeto = db.query(models.Projeto).filter(models.Projeto.id == id).first()
        
        # Caso a ID informada esteja presente no banco de dados
        if projeto != None:
            # Caso seja desejado atualizar o número de instalação
            if projeto_atualizado.numero_instalacao != None:
                projeto.numero_instalacao = projeto_atualizado.numero_instalacao
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o número do cliente
            if projeto_atualizado.numero_cliente != None:
                projeto.numero_cliente = projeto_atualizado.numero_cliente
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se é ligação nova
            if projeto_atualizado.ligacao_nova != None:
                projeto.ligacao_nova = projeto_atualizado.ligacao_nova
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se é aumento de carga
            if projeto_atualizado.aumento_carga != None:
                projeto.aumento_carga = projeto_atualizado.aumento_carga
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se é aumento de usina
            if projeto_atualizado.aumento_usina != None:
                projeto.aumento_usina = projeto_atualizado.aumento_usina
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se é agrupamento
            if projeto_atualizado.agrupamento != None:
                projeto.agrupamento = projeto_atualizado.agrupamento
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o número de fases
            if projeto_atualizado.n_fases != None:
                projeto.n_fases = projeto_atualizado.n_fases
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o disjuntor
            if projeto_atualizado.disjuntor != None:
                projeto.disjuntor = projeto_atualizado.disjuntor
                projeto.criado_por = usuario_atual.id
                
            # Caso seja desejado atualizar o novo número de fases
            if projeto_atualizado.novo_n_fases != None:
                projeto.novo_n_fases = projeto_atualizado.novo_n_fases
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o novo disjuntor
            if projeto_atualizado.novo_disjuntor != None:
                projeto.novo_disjuntor = projeto_atualizado.novo_disjuntor
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o número de fases do agrupamento
            if projeto_atualizado.n_fases_agrupamento != None:
                projeto.n_fases_agrupamento = projeto_atualizado.n_fases_agrupamento
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o disjuntor do agrupamento
            if projeto_atualizado.disjuntor_agrupamento != None:
                projeto.disjuntor_agrupamento = projeto_atualizado.disjuntor_agrupamento
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 1 é da usina anterior
            if projeto_atualizado.modulo_anterior_1 != None:
                projeto.modulo_anterior_1 = projeto_atualizado.modulo_anterior_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 1
            if projeto_atualizado.quantidade_modulo_1 != None:
                projeto.quantidade_modulo_1 = projeto_atualizado.quantidade_modulo_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 1
            if projeto_atualizado.modelo_modulo_1 != None:
                projeto.modelo_modulo_1 = projeto_atualizado.modelo_modulo_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 2 é da usina anterior
            if projeto_atualizado.modulo_anterior_2 != None:
                projeto.modulo_anterior_2 = projeto_atualizado.modulo_anterior_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 2
            if projeto_atualizado.quantidade_modulo_2 != None:
                projeto.quantidade_modulo_2 = projeto_atualizado.quantidade_modulo_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 2
            if projeto_atualizado.modelo_modulo_2 != None:
                projeto.modelo_modulo_2 = projeto_atualizado.modelo_modulo_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 2 é da usina anterior
            if projeto_atualizado.inversor_anterior_1 != None:
                projeto.inversor_anterior_1 = projeto_atualizado.inversor_anterior_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 2
            if projeto_atualizado.quantidade_inversor_1 != None:
                projeto.quantidade_inversor_1 = projeto_atualizado.quantidade_inversor_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 2
            if projeto_atualizado.modelo_inversor_1 != None:
                projeto.modelo_inversor_1 = projeto_atualizado.modelo_inversor_1
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 2 é da usina anterior
            if projeto_atualizado.inversor_anterior_2 != None:
                projeto.inversor_anterior_2 = projeto_atualizado.inversor_anterior_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 2
            if projeto_atualizado.quantidade_inversor_2 != None:
                projeto.quantidade_inversor_2 = projeto_atualizado.quantidade_inversor_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 2
            if projeto_atualizado.modelo_inversor_2 != None:
                projeto.modelo_inversor_2 = projeto_atualizado.modelo_inversor_2
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 2 é da usina anterior
            if projeto_atualizado.inversor_anterior_3 != None:
                projeto.inversor_anterior_3 = projeto_atualizado.inversor_anterior_3
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 2
            if projeto_atualizado.quantidade_inversor_3 != None:
                projeto.quantidade_inversor_3 = projeto_atualizado.quantidade_inversor_3
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 2
            if projeto_atualizado.modelo_inversor_3 != None:
                projeto.modelo_inversor_3 = projeto_atualizado.modelo_inversor_3
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar se o módulo 2 é da usina anterior
            if projeto_atualizado.inversor_anterior_4 != None:
                projeto.inversor_anterior_4 = projeto_atualizado.inversor_anterior_4
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar a quantidade de módulos 2
            if projeto_atualizado.quantidade_inversor_4 != None:
                projeto.quantidade_inversor_4 = projeto_atualizado.quantidade_inversor_4
                projeto.criado_por = usuario_atual.id

            # Caso seja desejado atualizar o modelo do módulo 2
            if projeto_atualizado.modelo_inversor_4 != None:
                projeto.modelo_inversor_4 = projeto_atualizado.modelo_inversor_4
                projeto.criado_por = usuario_atual.id

            # Determinando carga da instalação
            projeto.carga = utils.carga_instalacao(projeto.aumento_carga, projeto.n_fases, projeto.disjuntor, projeto.novo_n_fases, projeto.novo_disjuntor)
            
            # Determinando a quantidade total de módulos
            projeto.quantidade_modulos = projeto.quantidade_modulo_1 + projeto.quantidade_modulo_2
            
            # Determinando a potência total dos módulos anteriores
            projeto.potencia_modulos_anterior = utils.potencia_total_modulos_anterior(
                projeto.modulo_anterior_1, projeto.quantidade_modulo_1, projeto.potencia_modulo_1,
                projeto.modulo_anterior_2, projeto.quantidade_modulo_2, projeto.potencia_modulo_2
            )
            
            # Determinando a potência total dos módulos
            projeto.potencia_modulos = utils.potencia_total_modulos(
                projeto.quantidade_modulo_1, projeto.potencia_modulo_1,
                projeto.quantidade_modulo_2, projeto.potencia_modulo_2
            )
            
            modulo_1 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto.modelo_modulo_1).first()
            modulo_2 = models.Modulo()

            if projeto.modelo_modulo_2 != '':
                modulo_2 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto.modelo_modulo_2).first()

            # Determinando a área total ocupada pelos módulos
            if projeto.modelo_modulo_2 != '':
                projeto.area = utils.area_modulos(
                    projeto.quantidade_modulo_1, modulo_1.comprimento, modulo_1.largura,
                    projeto.quantidade_modulo_2, modulo_2.comprimento, modulo_2.largura
                )

            else:
                projeto.area = utils.area_modulos(
                    projeto.quantidade_modulo_1, modulo_1.comprimento, modulo_1.largura,
                    0, 0, 0
                )
            # Determinando a quantidade total de inversores
            projeto.quantidade_inversores = projeto.quantidade_inversor_1 + projeto.quantidade_inversor_2 + projeto.quantidade_inversor_3 + projeto.quantidade_inversor_4
            
            # Determinando a potência total dos módulos anteriores
            projeto.potencia_inversores_anterior = utils.potencia_total_inversores_anterior(
                projeto.inversor_anterior_1, projeto.quantidade_inversor_1, projeto.potencia_inversor_1,
                projeto.inversor_anterior_2, projeto.quantidade_inversor_2, projeto.potencia_inversor_2,
                projeto.inversor_anterior_3, projeto.quantidade_inversor_3, projeto.potencia_inversor_3,
                projeto.inversor_anterior_4, projeto.quantidade_inversor_4, projeto.potencia_inversor_4
            )

            # Determinando a potência total dos módulos anteriores
            projeto.potencia_inversores = utils.potencia_total_inversores(
                projeto.quantidade_inversor_1, projeto.potencia_inversor_1,
                projeto.quantidade_inversor_2, projeto.potencia_inversor_2,
                projeto.quantidade_inversor_3, projeto.potencia_inversor_3,
                projeto.quantidade_inversor_4, projeto.potencia_inversor_4
            )
            
            # Determinando a geração mensal
            cidade = db.query(models.Instalacao).filter(models.Instalacao.numero_instalacao == projeto.numero_instalacao).first().cidade
            projeto.geracao_1 = utils.geracao(projeto.potencia_modulos, cidade, 1)
            projeto.geracao_2 = utils.geracao(projeto.potencia_modulos, cidade, 2)
            projeto.geracao_3 = utils.geracao(projeto.potencia_modulos, cidade, 3)
            projeto.geracao_4 = utils.geracao(projeto.potencia_modulos, cidade, 4)
            projeto.geracao_5 = utils.geracao(projeto.potencia_modulos, cidade, 5)
            projeto.geracao_6 = utils.geracao(projeto.potencia_modulos, cidade, 6)
            projeto.geracao_7 = utils.geracao(projeto.potencia_modulos, cidade, 7)
            projeto.geracao_8 = utils.geracao(projeto.potencia_modulos, cidade, 8)
            projeto.geracao_9 = utils.geracao(projeto.potencia_modulos, cidade, 9)
            projeto.geracao_10 = utils.geracao(projeto.potencia_modulos, cidade, 10)
            projeto.geracao_11 = utils.geracao(projeto.potencia_modulos, cidade, 11)
            projeto.geracao_12 = utils.geracao(projeto.potencia_modulos, cidade, 12)

            db.commit() # Atualizando as informações no banco de dados
            return projeto
        
        # Caso a ID informada não esteja presente no banco de dados
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Projeto com a ID {id} não existe')

    # Caso o usuário atual não esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de atualizar um projeto')
    
# Função para deletar um projeto
@router.delete('/projects/{id}', status_code=status.HTTP_204_NO_CONTENT)
def delete_project(id: int, db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        projeto = db.query(models.Projeto).filter(models.Projeto.id == id)
        
        # Caso a ID informada esteja presente no banco de dados
        try:
            if projeto.first().id != None:
                projeto.delete(synchronize_session=False)
                db.commit()

        # Caso a ID informada não esteja presente no banco de dados
        except:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Projeto com a ID {id} não existe')
    
    # Caso o usuário atual esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de deletar um projeto')
    
# Função para imprimir os dados de um projeto
@router.get('/projects/{id}/print')
def print_project(id: int, db: Session = Depends(get_db), usuario_atual: models.Usuario = Depends(oauth2.determinar_usuario_atual)):
    # Caso o usuário atual esteja validado
    if usuario_atual != None:
        warnings.simplefilter("ignore")
        projeto = db.query(models.Projeto).filter(models.Projeto.id == id).first()

        # Caso a ID informada esteja presente no banco de dados
        if projeto != None:
            workbook = load_workbook('C://Users//Henrique Castro//Documents//entrada-dados//workbooks//04 - Formulario de Solicitacao de Acesso.xlsx')
            sheet = workbook.get_sheet_by_name('Formulário')
            
            projeto = db.query(models.Projeto).filter(models.Projeto.id == id).first()
            cliente = db.query(models.Cliente).filter(models.Cliente.numero_cliente == projeto.numero_cliente).first()
            instalacao = db.query(models.Instalacao).filter(models.Instalacao.numero_instalacao == projeto.numero_instalacao).first()
            modulo_1 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto.modelo_modulo_1).first()
            modulo_2 = models.Modulo()

            if projeto.modelo_modulo_2 != "":
                modulo_2 = db.query(models.Modulo).filter(models.Modulo.modelo == projeto.modelo_modulo_2).first()

            inversor_1 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto.modelo_inversor_1).first()
            inversor_2 = models.Inversor()
            inversor_3 = models.Inversor()
            inversor_4 = models.Inversor()

            if projeto.modelo_inversor_2 != "":
                inversor_2 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto.modelo_inversor_2).first()

            if projeto.modelo_inversor_3 != "":
                inversor_3 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto.modelo_inversor_3).first()

            if projeto.modelo_inversor_4 != "":
                inversor_4 = db.query(models.Inversor).filter(models.Inversor.modelo == projeto.modelo_inversor_4).first()

            fases: str = ""
            novo_fases: str = ""
                
            if projeto.n_fases == "Monopolar":
                fases = "Monofásico"

            elif projeto.n_fases == "Bipolar":
                fases = "Bifásico"

            elif projeto.n_fases == "Tripolar":
                fases = "Trifásico"

            if projeto.novo_n_fases == "Monopolar":
                novo_fases = "Monofásico"

            elif projeto.novo_n_fases == "Bipolar":
                novo_fases = "Bifásico"

            elif projeto.novo_n_fases == "Tripolar":
                novo_fases = "Trifásico"
            
            # Criando diretório para o projeto
            numero_inicial: int = 1
            numero_projeto = numero_inicial + projeto.id
            nome_completo = cliente.nome.split(" ")
            nome_pasta = ""

            for nome in nome_completo:
                nome_pasta += nome[0] + nome[1:].lower() + " "

            nome_completo = nome_pasta[:-1]

            if numero_projeto < 10:
                diretorio = f"G:/Meu Drive/Projetos fotovoltaicos/Projetos elaborados/00{numero_projeto} - {nome_completo}"
                diretorio_alterado = f"G:\\Meu Drive\\Projetos fotovoltaicos\\Projetos elaborados\\00{numero_projeto} - {nome_completo}"

            elif numero_projeto >= 10 and numero_projeto < 100:
                diretorio = f"G:/Meu Drive/Projetos fotovoltaicos/Projetos elaborados/0{numero_projeto} - {nome_completo}"
                diretorio_alterado = f"G:\\Meu Drive\\Projetos fotovoltaicos\\Projetos elaborados\\0{numero_projeto} - {nome_completo}"

            elif numero_projeto >= 100:
                diretorio = f"G:/Meu Drive/Projetos fotovoltaicos/Projetos elaborados/{numero_projeto} - {nome_completo}"
                diretorio_alterado = f"G:\\Meu Drive\\Projetos fotovoltaicos\\Projetos elaborados\\{numero_projeto} - {nome_completo}"

            
            if not os.path.exists(diretorio):
                os.makedirs(diretorio)
            
            # Preenchendo dados da identificação da unidade consumidora
            sheet["I12"] = projeto.numero_cliente

            if projeto.ligacao_nova == False:
                sheet["AG12"] = projeto.numero_instalacao

            sheet["N14"] = cliente.nome
            sheet["E16"] = "B"
            sheet["K16"] = instalacao.classificacao
            sheet["AC16"] = cliente.cpf
            sheet["G18"] = instalacao.logradouro
            sheet["AI18"] = instalacao.numero_predial
            sheet["AR18"] = instalacao.complemento
            sheet["E20"] = instalacao.bairro
            sheet["T20"] = instalacao.cidade
            sheet["AN20"] = "MG"
            sheet["AS20"] = instalacao.cep
            sheet["F22"] = 3798262232
            sheet["O22"] = 37998262232
            sheet["Y22"] = "engenharia@apollofotovoltaica.com.br"

            # Preenchendo dados da unidade consumidora
            sheet["V27"] = 23
            sheet["AC27"] = instalacao.longitude
            sheet["AL27"] = instalacao.latitude

            potencia_usina_antiga: int = 0

            if projeto.potencia_modulos_anterior < projeto.potencia_inversores_anterior:
                potencia_usina_antiga = projeto.potencia_modulos_anterior

            else:
                potencia_usina_antiga = projeto.potencia_inversores_anterior

            if projeto.ligacao_nova == True:
                sheet["I39"] = "Ligação de Nova Unidade Consumidora COM Geração Distribuída"

            potencia_usina: int = 0

            if projeto.potencia_modulos < projeto.potencia_inversores:
                potencia_usina = projeto.potencia_modulos

            else:
                potencia_usina = projeto.potencia_inversores

            if projeto.ligacao_nova == True:
                sheet["I39"] = "Ligação de Nova Unidade Consumidora COM Geração Distribuída"

            elif projeto.aumento_usina == True:
                sheet["B7"] = ""
                sheet["I39"] = "GD Existente COM Alteração de Potência Ativa Instalada Total"
                sheet["AU39"] = potencia_usina_antiga
                sheet["AU41"] = potencia_usina

            else:
                if projeto.aumento_carga == False:
                    sheet["I39"] = "Conexão de GD em Unidade Consumidora Existente SEM Alteração de Potência Disponibilizada"

                elif projeto.aumento_carga == True:
                    sheet["I39"] = "Conexão de GD em Unidade Consumidora Existente COM Alteração de Potência Disponibilizada"

            sheet["AB45"] = projeto.n_fases
            sheet["AH45"] = projeto.disjuntor

            if projeto.aumento_carga == True:
                sheet["AB47"] = projeto.novo_n_fases
                sheet["AH47"] = projeto.novo_disjuntor
                sheet["R54"] = "Sim"

                if projeto.agrupamento == True:
                    sheet["AB49"] = projeto.n_fases_agrupamento
                    sheet["AH49"] = projeto.disjuntor_agrupamento
                    sheet["AT49"] = 1

            else:
                sheet["R54"] = "Não"

                if projeto.agrupamento == True:
                    sheet["AB47"] = projeto.n_fases_agrupamento
                    sheet["AH47"] = projeto.disjuntor_agrupamento
                    sheet["AT47"] = 1

            if projeto.agrupamento == False:
                sheet["I41"] = "Edificação Individual"

            else:
                sheet["I41"] = "Agrupamento"

            if instalacao.classificacao == "Rural":
                sheet["L52"] = "120/240"

            else:
                sheet["L52"] = "127/220"

            sheet["X52"] = "Aéreo"

            # Preenchendo dados da geração
            sheet["L80"] = "Solar"
            sheet["H82"] = "Empregando conversor eletrônico/inversor"
            sheet["L84"] = "Compensação local"
            sheet["AS84"] = 1

            sheet["P92"] = projeto.potencia_modulos
            sheet["P94"] = projeto.potencia_inversores
            sheet["P96"] = projeto.area
            sheet["P98"] = projeto.quantidade_modulos

            fabricante_modulos = modulo_1.fabricante
            modelo_modulos = projeto.modelo_modulo_1
            
            if modulo_2.modelo != None:
                fabricante_modulos = fabricante_modulos + " / " + modulo_2.fabricante
                modelo_modulos = modelo_modulos + " / " + projeto.modelo_modulo_2

            sheet["P100"] = modelo_modulos
            sheet["P102"] = fabricante_modulos
            sheet["P104"] = projeto.quantidade_inversores

            fabricante_inversores = inversor_1.fabricante
            modelo_inversores = projeto.modelo_inversor_1
            
            if inversor_2.modelo != None:
                fabricante_inversores = fabricante_inversores + " / " + inversor_2.fabricante
                modelo_inversores = modelo_inversores + " / " + projeto.modelo_inversor_2

            if inversor_3.modelo != None:
                fabricante_inversores = fabricante_inversores + " / " + inversor_3.fabricante
                modelo_inversores = modelo_inversores + " / " + projeto.modelo_inversor_3

            if inversor_4.modelo != None:
                fabricante_inversores = fabricante_inversores + " / " + inversor_4.fabricante
                modelo_inversores = modelo_inversores + " / " + projeto.modelo_inversor_4

            sheet["P106"] = modelo_inversores
            sheet["P108"] = fabricante_inversores

            # Preenchendo dados do solicitante
            sheet["Q194"] = "André Luís Queiroz Nogueira"
            sheet["O196"] = "Rua Agripino Lima, 68 Ap. 202, Centro, Itaúna - MG"
            sheet["F200"] = 3798262232
            sheet["O200"] = 37998262232
            sheet["Y200"] = "engenharia@apollofotovoltaica.com.br"

            dia_atual = date.today().strftime("%d/%m/%Y")
            sheet["C210"] = instalacao.cidade + ", " + dia_atual

            assinatura = Image('C://Users//Henrique Castro//Documents//entrada-dados//workbooks//assinatura.png')
            sheet.add_image(assinatura, "V210")
            
            workbook.save(f'{diretorio}/04 - Formulario de Solicitacao de Acesso - {nome_completo}.xlsx')
            workbook.close()

            if projeto.aumento_carga == False:
                memorial = docx.Document("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\02 - Memorial Descritivo.docx")

                # Alterando informações da capa do memorial
                memorial.paragraphs[17].text = memorial.paragraphs[17].text.replace("XXXX", instalacao.classificacao)
                memorial.paragraphs[31].text = memorial.paragraphs[31].text.replace("XXXX", instalacao.cidade)
                memorial.paragraphs[32].text = memorial.paragraphs[32].text.replace("XXXX", dia_atual)
                
                # Alterando informações da seção objetivo
                memorial.paragraphs[55].text = memorial.paragraphs[55].text.replace("XXXX", instalacao.cidade)

                # Alterando informações da seção dados da unidade consumidora
                memorial.paragraphs[58].text = memorial.paragraphs[58].text.replace("XXXX", cliente.nome)
                memorial.paragraphs[59].text = memorial.paragraphs[59].text.replace("XXXX", instalacao.logradouro + ", " + str(instalacao.numero_predial) + " " + instalacao.complemento + ", " + instalacao.bairro + ", " + instalacao.cidade + " - MG")
                memorial.paragraphs[60].text = memorial.paragraphs[60].text.replace("XXXX", cliente.cpf)
                
                if projeto.ligacao_nova == False:
                    memorial.paragraphs[61].text = memorial.paragraphs[61].text.replace("XXXX", str(projeto.numero_instalacao))
                    
                else:
                    memorial.paragraphs[61].text = memorial.paragraphs[61].text.replace("XXXX", "")
                    memorial.paragraphs[65].text = memorial.paragraphs[65].text.replace("atual", "solicitado")

                memorial.paragraphs[62].text = memorial.paragraphs[62].text.replace("XXXX", str(projeto.numero_cliente))
                memorial.paragraphs[63].text = memorial.paragraphs[63].text.replace("XXXX", instalacao.classificacao)
                
                fases: str

                if projeto.n_fases == "Monopolar":
                    fases = "Monofásico"

                elif projeto.n_fases == "Bipolar":
                    fases = "Bifásico"

                elif projeto.n_fases == "Tripolar":
                    fases = "Trifásico"

                memorial.paragraphs[64].text = memorial.paragraphs[64].text.replace("XXXX", instalacao.classificacao + " " + fases)
                memorial.paragraphs[65].text = memorial.paragraphs[65].text.replace("XXXX", projeto.n_fases + " de " + str(projeto.disjuntor) + " A")
                memorial.paragraphs[67].text = memorial.paragraphs[67].text.replace("XXXX", str(projeto.area))
                memorial.paragraphs[67].text = memorial.paragraphs[67].text.replace(".", ",")
                memorial.paragraphs[71].text = memorial.paragraphs[71].text.replace("XXXX", str(instalacao.longitude))
                memorial.paragraphs[72].text = memorial.paragraphs[72].text.replace("XXXX", str(instalacao.latitude))
                
                memorial.paragraphs[85].text = ""
                paragrafo_imagem = memorial.paragraphs[85]
                run_imagem = paragrafo_imagem.add_run()
                run_imagem.add_picture("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\foto-satelite.jpg", height=Inches(3.54331))

                memorial.paragraphs[87].text = ""
                paragrafo_imagem = memorial.paragraphs[87]
                run_imagem = paragrafo_imagem.add_run()
                run_imagem.add_picture("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\foto-disjuntor.jpeg", height=Inches(3.54331))
                
                paragrafo_atual: int = 122
                incremento: int = 0

                # Alterando informações da seção descrição geral da geração distribuída
                if projeto.modelo_modulo_2 != "":
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[102].text.replace("XXXX", str(modulo_2.comprimento) + " x " + str(modulo_2.largura) + " x " + str(modulo_2.espessura) + " mm"))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[101].text.replace("XXXX", modulo_2.temperatura_nominal))
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[100].text.replace("XXXX", str(modulo_2.eficiencia)))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[99].text.replace("XXXX", str(modulo_2.voc) + " V"))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[98].text.replace("XXXX", str(modulo_2.vmp) + " V"))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[97].text.replace("XXXX", str(modulo_2.isc) + " A"))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[96].text.replace("XXXX", str(modulo_2.imp) + " A"))
                    memorial.paragraphs[103].text = memorial.paragraphs[103].text.replace(".", ",")
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[87].text.replace("XXXX", str(modulo_2.potencia) + " Wp"))
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[86].text.replace("XXXX", modulo_2.tipo))
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[85].text.replace("XXXX", modulo_2.fabricante))
                    memorial.paragraphs[103].insert_paragraph_before(memorial.paragraphs[84].text.replace("XXXX", modulo_2.modelo))

                    memorial.paragraphs[103].style = "Tópicos"
                    memorial.paragraphs[100].style = "Tópicos"
                    memorial.paragraphs[101].style = "Tópicos"
                    memorial.paragraphs[102].style = "Tópicos"
                    memorial.paragraphs[103].style = "Tópicos"
                    memorial.paragraphs[104].style = "Tópicos"
                    memorial.paragraphs[105].style = "Tópicos"
                    memorial.paragraphs[106].style = "Tópicos"
                    memorial.paragraphs[107].style = "Tópicos"
                    memorial.paragraphs[108].style = "Tópicos"
                    memorial.paragraphs[109].style = "Tópicos"
                    
                    memorial.paragraphs[103].insert_paragraph_before()
                    paragrafo_atual += 12
                
                memorial.paragraphs[92].text = memorial.paragraphs[92].text.replace("XXXX", modulo_1.modelo)
                memorial.paragraphs[93].text = memorial.paragraphs[93].text.replace("XXXX", modulo_1.fabricante)
                memorial.paragraphs[94].text = memorial.paragraphs[94].text.replace("XXXX", modulo_1.tipo)
                memorial.paragraphs[95].text = memorial.paragraphs[95].text.replace("XXXX", str(modulo_1.potencia) + " Wp")
                memorial.paragraphs[96].text = memorial.paragraphs[96].text.replace("XXXX", str(modulo_1.imp) + " A")
                memorial.paragraphs[96].text = memorial.paragraphs[96].text.replace(".", ",")
                memorial.paragraphs[97].text = memorial.paragraphs[97].text.replace("XXXX", str(modulo_1.isc) + " A")
                memorial.paragraphs[97].text = memorial.paragraphs[97].text.replace(".", ",")
                memorial.paragraphs[98].text = memorial.paragraphs[98].text.replace("XXXX", str(modulo_1.vmp) + " V")
                memorial.paragraphs[98].text = memorial.paragraphs[98].text.replace(".", ",")
                memorial.paragraphs[99].text = memorial.paragraphs[99].text.replace("XXXX", str(modulo_1.voc) + " V")
                memorial.paragraphs[99].text = memorial.paragraphs[99].text.replace(".", ",")
                memorial.paragraphs[100].text = memorial.paragraphs[100].text.replace("XXXX", str(modulo_1.eficiencia) + "%")
                memorial.paragraphs[100].text = memorial.paragraphs[100].text.replace(".", ",")
                memorial.paragraphs[101].text = memorial.paragraphs[101].text.replace("XXXX", modulo_1.temperatura_nominal)
                memorial.paragraphs[102].text = memorial.paragraphs[102].text.replace("XXXX", str(modulo_1.comprimento) + " x " + str(modulo_1.largura) + " x " + str(modulo_1.espessura) + " mm")
                memorial.paragraphs[102].text = memorial.paragraphs[102].text.replace(".", ",")

                memorial.paragraphs[paragrafo_atual - 17].text = memorial.paragraphs[paragrafo_atual - 17].text.replace("XXXX", str(projeto.quantidade_inversores))
                memorial.paragraphs[paragrafo_atual - 16].text = memorial.paragraphs[paragrafo_atual - 16].text.replace("XXXX", str(projeto.potencia_inversores) + " kW")
                memorial.paragraphs[paragrafo_atual - 16].text = memorial.paragraphs[paragrafo_atual - 16].text.replace(".", ",")

                if projeto.modelo_inversor_4 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_4.comprimento) + " x " + str(inversor_4.largura) + " x " + str(inversor_4.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_4.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_4.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_4.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_4.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_4.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_4.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_4.v_min_mppt) + "-" + str(inversor_4.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_4.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_4.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_4.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_4.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_4.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_4.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()

                    incremento += 15

                if projeto.modelo_inversor_3 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_3.comprimento) + " x " + str(inversor_3.largura) + " x " + str(inversor_3.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_3.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_3.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_3.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_3.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_3.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_3.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_3.v_min_mppt) + "-" + str(inversor_3.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_3.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_3.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_3.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_3.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_3.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_3.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()

                    incremento += 15

                if projeto.modelo_inversor_2 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_2.comprimento) + " x " + str(inversor_2.largura) + " x " + str(inversor_2.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_2.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_2.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_2.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_2.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_2.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_2.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_2.v_min_mppt) + "-" + str(inversor_2.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_2.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_2.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_2.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_2.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_2.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_2.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()
                    
                    incremento += 15

                memorial.paragraphs[paragrafo_atual - 14].text = memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_1.modelo)
                memorial.paragraphs[paragrafo_atual - 13].text = memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_1.fabricante)
                memorial.paragraphs[paragrafo_atual - 12].text = memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_1.potencia / 1000) + " kW")
                memorial.paragraphs[paragrafo_atual - 12].text = memorial.paragraphs[paragrafo_atual - 12].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 11].text = memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_1.overload / 1000) + " kW")
                memorial.paragraphs[paragrafo_atual - 11].text = memorial.paragraphs[paragrafo_atual - 11].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 10].text = memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_1.imp) + " A")
                memorial.paragraphs[paragrafo_atual - 10].text = memorial.paragraphs[paragrafo_atual - 10].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 9].text = memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_1.isc) + " A")
                memorial.paragraphs[paragrafo_atual - 9].text = memorial.paragraphs[paragrafo_atual - 9].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 8].text = memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_1.v_min_mppt) + "-" + str(inversor_1.v_max_mppt) + " V")
                memorial.paragraphs[paragrafo_atual - 7].text = memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_1.v_max) + " V")
                memorial.paragraphs[paragrafo_atual - 6].text = memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_1.n_mppt))
                memorial.paragraphs[paragrafo_atual - 5].text = memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_1.n_entrada))
                memorial.paragraphs[paragrafo_atual - 4].text = memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_1.v_saida) + " V")
                memorial.paragraphs[paragrafo_atual - 3].text = memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_1.i_saida) + " A")
                memorial.paragraphs[paragrafo_atual - 3].text = memorial.paragraphs[paragrafo_atual - 3].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 2].text = memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_1.eficiencia) + "%")
                memorial.paragraphs[paragrafo_atual - 2].text = memorial.paragraphs[paragrafo_atual - 2].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 1].text = memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_1.comprimento) + " x " + str(inversor_1.largura) + " x " + str(inversor_1.espessura) + " mm")
                memorial.paragraphs[paragrafo_atual - 1].text = memorial.paragraphs[paragrafo_atual - 1].text.replace(".", ",")

                paragrafo_atual += incremento

                # Alterando informações da seção previsão da produção de energia
                memorial.paragraphs[paragrafo_atual + 43].text = memorial.paragraphs[paragrafo_atual + 43].text.replace("XXXX", str(projeto.geracao_1) + " kWh")
                memorial.paragraphs[paragrafo_atual + 43].text = memorial.paragraphs[paragrafo_atual + 43].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 44].text = memorial.paragraphs[paragrafo_atual + 44].text.replace("XXXX", str(projeto.geracao_2) + " kWh")
                memorial.paragraphs[paragrafo_atual + 44].text = memorial.paragraphs[paragrafo_atual + 44].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 45].text = memorial.paragraphs[paragrafo_atual + 45].text.replace("XXXX", str(projeto.geracao_3) + " kWh")
                memorial.paragraphs[paragrafo_atual + 45].text = memorial.paragraphs[paragrafo_atual + 45].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 46].text = memorial.paragraphs[paragrafo_atual + 46].text.replace("XXXX", str(projeto.geracao_4) + " kWh")
                memorial.paragraphs[paragrafo_atual + 46].text = memorial.paragraphs[paragrafo_atual + 46].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 47].text = memorial.paragraphs[paragrafo_atual + 47].text.replace("XXXX", str(projeto.geracao_5) + " kWh")
                memorial.paragraphs[paragrafo_atual + 47].text = memorial.paragraphs[paragrafo_atual + 47].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 48].text = memorial.paragraphs[paragrafo_atual + 48].text.replace("XXXX", str(projeto.geracao_6) + " kWh")
                memorial.paragraphs[paragrafo_atual + 48].text = memorial.paragraphs[paragrafo_atual + 48].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 49].text = memorial.paragraphs[paragrafo_atual + 49].text.replace("XXXX", str(projeto.geracao_7) + " kWh")
                memorial.paragraphs[paragrafo_atual + 49].text = memorial.paragraphs[paragrafo_atual + 49].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 50].text = memorial.paragraphs[paragrafo_atual + 50].text.replace("XXXX", str(projeto.geracao_8) + " kWh")
                memorial.paragraphs[paragrafo_atual + 50].text = memorial.paragraphs[paragrafo_atual + 50].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 51].text = memorial.paragraphs[paragrafo_atual + 51].text.replace("XXXX", str(projeto.geracao_9) + " kWh")
                memorial.paragraphs[paragrafo_atual + 51].text = memorial.paragraphs[paragrafo_atual + 51].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 52].text = memorial.paragraphs[paragrafo_atual + 52].text.replace("XXXX", str(projeto.geracao_10) + " kWh")
                memorial.paragraphs[paragrafo_atual + 52].text = memorial.paragraphs[paragrafo_atual + 52].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 53].text = memorial.paragraphs[paragrafo_atual + 53].text.replace("XXXX", str(projeto.geracao_11) + " kWh")
                memorial.paragraphs[paragrafo_atual + 53].text = memorial.paragraphs[paragrafo_atual + 53].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 54].text = memorial.paragraphs[paragrafo_atual + 54].text.replace("XXXX", str(projeto.geracao_12) + " kWh")
                memorial.paragraphs[paragrafo_atual + 54].text = memorial.paragraphs[paragrafo_atual + 54].text.replace(".", ",")
                
                memorial.save(f'{diretorio_alterado}\\02 - Memorial Descritivo - {nome_completo}.docx')

            elif projeto.aumento_carga == True:
                memorial = docx.Document("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\02 - Memorial Descritivo - Com Aumento.docx")

                # Alterando informações da capa do memorial
                memorial.paragraphs[17].text = memorial.paragraphs[17].text.replace("XXXX", instalacao.classificacao)
                memorial.paragraphs[29].text = memorial.paragraphs[29].text.replace("XXXX", instalacao.cidade)
                memorial.paragraphs[30].text = memorial.paragraphs[30].text.replace("XXXX", dia_atual)
                
                memorial.paragraphs[52].text = memorial.paragraphs[52].text.replace("XXXX", instalacao.cidade)

                # Alterando informações da seção dados da unidade consumidora
                memorial.paragraphs[55].text = memorial.paragraphs[55].text.replace("XXXX", cliente.nome)
                memorial.paragraphs[56].text = memorial.paragraphs[56].text.replace("XXXX", instalacao.logradouro + ", " + str(instalacao.numero_predial) + " " + instalacao.complemento + ", " + instalacao.bairro + ", " + instalacao.cidade + " - MG")
                memorial.paragraphs[57].text = memorial.paragraphs[57].text.replace("XXXX", cliente.cpf)
                memorial.paragraphs[58].text = memorial.paragraphs[58].text.replace("XXXX", str(projeto.numero_instalacao))
                memorial.paragraphs[59].text = memorial.paragraphs[59].text.replace("XXXX", str(projeto.numero_cliente))
                memorial.paragraphs[60].text = memorial.paragraphs[60].text.replace("XXXX", instalacao.classificacao)
                
                memorial.paragraphs[61].text = memorial.paragraphs[61].text.replace("XXXX", instalacao.classificacao + " " + novo_fases)
                memorial.paragraphs[62].text = memorial.paragraphs[62].text.replace("XXXX", projeto.n_fases + " de " + str(projeto.disjuntor) + " A")
                memorial.paragraphs[63].text = memorial.paragraphs[63].text.replace("XXXX", projeto.novo_n_fases + " de " + str(projeto.novo_disjuntor) + " A")
                memorial.paragraphs[65].text = memorial.paragraphs[65].text.replace("XXXX", str(projeto.area))
                memorial.paragraphs[65].text = memorial.paragraphs[65].text.replace(".", ",")
                memorial.paragraphs[69].text = memorial.paragraphs[69].text.replace("XXXX", str(instalacao.longitude))
                memorial.paragraphs[70].text = memorial.paragraphs[70].text.replace("XXXX", str(instalacao.latitude))
                
                memorial.paragraphs[82].text = ""
                paragrafo_imagem = memorial.paragraphs[82]
                run_imagem = paragrafo_imagem.add_run()
                run_imagem.add_picture("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\foto-satelite.jpg", height=Inches(3.14961))

                memorial.paragraphs[84].text = ""
                paragrafo_imagem = memorial.paragraphs[84]
                run_imagem = paragrafo_imagem.add_run()
                run_imagem.add_picture("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\foto-disjuntor.jpeg", height=Inches(3.54331))

                memorial.paragraphs[88].text = ""
                paragrafo_imagem = memorial.paragraphs[88]
                run_imagem = paragrafo_imagem.add_run()
                run_imagem.add_picture("C:\\Users\\Henrique Castro\\Documents\\entrada-dados\\workbooks\\foto-aumento.jpg", height=Inches(3.54331))
                
                memorial.paragraphs[94].text = memorial.paragraphs[94].text.replace("XXXX", str(instalacao.latitude + 2))
                memorial.paragraphs[95].text = memorial.paragraphs[95].text.replace("XXXX", str(instalacao.longitude + 2))
                memorial.paragraphs[98].text = memorial.paragraphs[98].text.replace("XXXX", str(instalacao.latitude))
                memorial.paragraphs[99].text = memorial.paragraphs[99].text.replace("XXXX", str(instalacao.longitude))

                paragrafo_atual: int = 141
                incremento: int = 0

                # Alterando informações da seção descrição geral da geração distribuída
                if projeto.modelo_modulo_2 != "":
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[121].text.replace("XXXX", str(modulo_2.comprimento) + " x " + str(modulo_2.largura) + " x " + str(modulo_2.espessura) + " mm"))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[120].text.replace("XXXX", modulo_2.temperatura_nominal))
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[119].text.replace("XXXX", str(modulo_2.eficiencia)))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[118].text.replace("XXXX", str(modulo_2.voc) + " V"))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[117].text.replace("XXXX", str(modulo_2.vmp) + " V"))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[116].text.replace("XXXX", str(modulo_2.isc) + " A"))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[115].text.replace("XXXX", str(modulo_2.imp) + " A"))
                    memorial.paragraphs[122].text = memorial.paragraphs[122].text.replace(".", ",")
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[114].text.replace("XXXX", str(modulo_2.potencia) + " Wp"))
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[113].text.replace("XXXX", modulo_2.tipo))
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[112].text.replace("XXXX", modulo_2.fabricante))
                    memorial.paragraphs[122].insert_paragraph_before(memorial.paragraphs[111].text.replace("XXXX", modulo_2.modelo))

                    memorial.paragraphs[122].style = "Tópicos"
                    memorial.paragraphs[123].style = "Tópicos"
                    memorial.paragraphs[124].style = "Tópicos"
                    memorial.paragraphs[125].style = "Tópicos"
                    memorial.paragraphs[126].style = "Tópicos"
                    memorial.paragraphs[127].style = "Tópicos"
                    memorial.paragraphs[128].style = "Tópicos"
                    memorial.paragraphs[129].style = "Tópicos"
                    memorial.paragraphs[130].style = "Tópicos"
                    memorial.paragraphs[131].style = "Tópicos"
                    memorial.paragraphs[132].style = "Tópicos"
                    
                    memorial.paragraphs[122].insert_paragraph_before()
                    paragrafo_atual += 12
                
                memorial.paragraphs[111].text = memorial.paragraphs[111].text.replace("XXXX", modulo_1.modelo)
                memorial.paragraphs[112].text = memorial.paragraphs[112].text.replace("XXXX", modulo_1.fabricante)
                memorial.paragraphs[113].text = memorial.paragraphs[113].text.replace("XXXX", modulo_1.tipo)
                memorial.paragraphs[114].text = memorial.paragraphs[114].text.replace("XXXX", str(modulo_1.potencia) + " Wp")
                memorial.paragraphs[115].text = memorial.paragraphs[115].text.replace("XXXX", str(modulo_1.imp) + " A")
                memorial.paragraphs[115].text = memorial.paragraphs[115].text.replace(".", ",")
                memorial.paragraphs[116].text = memorial.paragraphs[116].text.replace("XXXX", str(modulo_1.isc) + " A")
                memorial.paragraphs[116].text = memorial.paragraphs[116].text.replace(".", ",")
                memorial.paragraphs[117].text = memorial.paragraphs[117].text.replace("XXXX", str(modulo_1.vmp) + " V")
                memorial.paragraphs[117].text = memorial.paragraphs[117].text.replace(".", ",")
                memorial.paragraphs[118].text = memorial.paragraphs[118].text.replace("XXXX", str(modulo_1.voc) + " V")
                memorial.paragraphs[118].text = memorial.paragraphs[118].text.replace(".", ",")
                memorial.paragraphs[119].text = memorial.paragraphs[119].text.replace("XXXX", str(modulo_1.eficiencia) + "%")
                memorial.paragraphs[119].text = memorial.paragraphs[119].text.replace(".", ",")
                memorial.paragraphs[120].text = memorial.paragraphs[120].text.replace("XXXX", modulo_1.temperatura_nominal)
                memorial.paragraphs[121].text = memorial.paragraphs[121].text.replace("XXXX", str(modulo_1.comprimento) + " x " + str(modulo_1.largura) + " x " + str(modulo_1.espessura) + " mm")
                memorial.paragraphs[121].text = memorial.paragraphs[121].text.replace(".", ",")

                memorial.paragraphs[paragrafo_atual - 17].text = memorial.paragraphs[paragrafo_atual - 17].text.replace("XXXX", str(projeto.quantidade_inversores))
                memorial.paragraphs[paragrafo_atual - 16].text = memorial.paragraphs[paragrafo_atual - 16].text.replace("XXXX", str(projeto.potencia_inversores) + " kW")
                memorial.paragraphs[paragrafo_atual - 16].text = memorial.paragraphs[paragrafo_atual - 16].text.replace(".", ",")

                if projeto.modelo_inversor_4 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_4.comprimento) + " x " + str(inversor_4.largura) + " x " + str(inversor_4.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_4.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_4.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_4.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_4.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_4.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_4.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_4.v_min_mppt) + "-" + str(inversor_4.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_4.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_4.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_4.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_4.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_4.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_4.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()
                    incremento += 15

                if projeto.modelo_inversor_3 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_3.comprimento) + " x " + str(inversor_3.largura) + " x " + str(inversor_3.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_3.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_3.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_3.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_3.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_3.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_3.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_3.v_min_mppt) + "-" + str(inversor_3.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_3.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_3.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_3.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_3.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_3.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_3.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()
                    incremento += 15

                if projeto.modelo_inversor_2 != "":
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_2.comprimento) + " x " + str(inversor_2.largura) + " x " + str(inversor_2.espessura) + " mm"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_2.eficiencia) + "%"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_2.i_saida) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_2.v_saida) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_2.n_entrada)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_2.n_mppt)))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_2.v_max) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_2.v_min_mppt) + "-" + str(inversor_2.v_max_mppt) + " V"))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_2.isc) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_2.imp) + " A"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_2.overload / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_2.potencia / 1000) + " kW"))
                    memorial.paragraphs[paragrafo_atual].text = memorial.paragraphs[paragrafo_atual].text.replace(".", ",")
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_2.fabricante))
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before(memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_2.modelo))

                    memorial.paragraphs[paragrafo_atual].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 1].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 2].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 3].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 4].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 5].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 6].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 7].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 8].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 9].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 10].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 11].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 12].style = "Subtópicos"
                    memorial.paragraphs[paragrafo_atual + 13].style = "Subtópicos"
                    
                    memorial.paragraphs[paragrafo_atual].insert_paragraph_before()
                    incremento += 15
                    
                memorial.paragraphs[paragrafo_atual - 14].text = memorial.paragraphs[paragrafo_atual - 14].text.replace("XXXX", inversor_1.modelo)
                memorial.paragraphs[paragrafo_atual - 13].text = memorial.paragraphs[paragrafo_atual - 13].text.replace("XXXX", inversor_1.fabricante)
                memorial.paragraphs[paragrafo_atual - 12].text = memorial.paragraphs[paragrafo_atual - 12].text.replace("XXXX", str(inversor_1.potencia / 1000) + " kW")
                memorial.paragraphs[paragrafo_atual - 12].text = memorial.paragraphs[paragrafo_atual - 12].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 11].text = memorial.paragraphs[paragrafo_atual - 11].text.replace("XXXX", str(inversor_1.overload / 1000) + " kW")
                memorial.paragraphs[paragrafo_atual - 11].text = memorial.paragraphs[paragrafo_atual - 11].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 10].text = memorial.paragraphs[paragrafo_atual - 10].text.replace("XXXX", str(inversor_1.imp) + " A")
                memorial.paragraphs[paragrafo_atual - 10].text = memorial.paragraphs[paragrafo_atual - 10].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 9].text = memorial.paragraphs[paragrafo_atual - 9].text.replace("XXXX", str(inversor_1.isc) + " A")
                memorial.paragraphs[paragrafo_atual - 9].text = memorial.paragraphs[paragrafo_atual - 9].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 8].text = memorial.paragraphs[paragrafo_atual - 8].text.replace("XXXX", str(inversor_1.v_min_mppt) + "-" + str(inversor_1.v_max_mppt) + " V")
                memorial.paragraphs[paragrafo_atual - 7].text = memorial.paragraphs[paragrafo_atual - 7].text.replace("XXXX", str(inversor_1.v_max) + " V")
                memorial.paragraphs[paragrafo_atual - 6].text = memorial.paragraphs[paragrafo_atual - 6].text.replace("XXXX", str(inversor_1.n_mppt))
                memorial.paragraphs[paragrafo_atual - 5].text = memorial.paragraphs[paragrafo_atual - 5].text.replace("XXXX", str(inversor_1.n_entrada))
                memorial.paragraphs[paragrafo_atual - 4].text = memorial.paragraphs[paragrafo_atual - 4].text.replace("XXXX", str(inversor_1.v_saida) + " V")
                memorial.paragraphs[paragrafo_atual - 3].text = memorial.paragraphs[paragrafo_atual - 3].text.replace("XXXX", str(inversor_1.i_saida) + " A")
                memorial.paragraphs[paragrafo_atual - 3].text = memorial.paragraphs[paragrafo_atual - 3].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 2].text = memorial.paragraphs[paragrafo_atual - 2].text.replace("XXXX", str(inversor_1.eficiencia) + "%")
                memorial.paragraphs[paragrafo_atual - 2].text = memorial.paragraphs[paragrafo_atual - 2].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual - 1].text = memorial.paragraphs[paragrafo_atual - 1].text.replace("XXXX", str(inversor_1.comprimento) + " x " + str(inversor_1.largura) + " x " + str(inversor_1.espessura) + " mm")
                memorial.paragraphs[paragrafo_atual - 1].text = memorial.paragraphs[paragrafo_atual - 1].text.replace(".", ",")
                
                paragrafo_atual += incremento

                # Alterando informações da seção previsão da produção de energia
                memorial.paragraphs[paragrafo_atual + 42].text = memorial.paragraphs[paragrafo_atual + 42].text.replace("XXXX", str(projeto.geracao_1) + " kWh")
                memorial.paragraphs[paragrafo_atual + 42].text = memorial.paragraphs[paragrafo_atual + 42].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 43].text = memorial.paragraphs[paragrafo_atual + 43].text.replace("XXXX", str(projeto.geracao_2) + " kWh")
                memorial.paragraphs[paragrafo_atual + 43].text = memorial.paragraphs[paragrafo_atual + 43].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 44].text = memorial.paragraphs[paragrafo_atual + 44].text.replace("XXXX", str(projeto.geracao_3) + " kWh")
                memorial.paragraphs[paragrafo_atual + 44].text = memorial.paragraphs[paragrafo_atual + 44].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 45].text = memorial.paragraphs[paragrafo_atual + 45].text.replace("XXXX", str(projeto.geracao_4) + " kWh")
                memorial.paragraphs[paragrafo_atual + 45].text = memorial.paragraphs[paragrafo_atual + 45].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 46].text = memorial.paragraphs[paragrafo_atual + 46].text.replace("XXXX", str(projeto.geracao_5) + " kWh")
                memorial.paragraphs[paragrafo_atual + 46].text = memorial.paragraphs[paragrafo_atual + 46].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 47].text = memorial.paragraphs[paragrafo_atual + 47].text.replace("XXXX", str(projeto.geracao_6) + " kWh")
                memorial.paragraphs[paragrafo_atual + 47].text = memorial.paragraphs[paragrafo_atual + 47].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 48].text = memorial.paragraphs[paragrafo_atual + 48].text.replace("XXXX", str(projeto.geracao_7) + " kWh")
                memorial.paragraphs[paragrafo_atual + 48].text = memorial.paragraphs[paragrafo_atual + 48].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 49].text = memorial.paragraphs[paragrafo_atual + 49].text.replace("XXXX", str(projeto.geracao_8) + " kWh")
                memorial.paragraphs[paragrafo_atual + 49].text = memorial.paragraphs[paragrafo_atual + 49].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 50].text = memorial.paragraphs[paragrafo_atual + 50].text.replace("XXXX", str(projeto.geracao_9) + " kWh")
                memorial.paragraphs[paragrafo_atual + 50].text = memorial.paragraphs[paragrafo_atual + 50].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 51].text = memorial.paragraphs[paragrafo_atual + 51].text.replace("XXXX", str(projeto.geracao_10) + " kWh")
                memorial.paragraphs[paragrafo_atual + 51].text = memorial.paragraphs[paragrafo_atual + 51].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 52].text = memorial.paragraphs[paragrafo_atual + 52].text.replace("XXXX", str(projeto.geracao_11) + " kWh")
                memorial.paragraphs[paragrafo_atual + 52].text = memorial.paragraphs[paragrafo_atual + 52].text.replace(".", ",")
                memorial.paragraphs[paragrafo_atual + 53].text = memorial.paragraphs[paragrafo_atual + 53].text.replace("XXXX", str(projeto.geracao_12) + " kWh")
                memorial.paragraphs[paragrafo_atual + 53].text = memorial.paragraphs[paragrafo_atual + 53].text.replace(".", ",")
                
                memorial.save(f"{diretorio_alterado}\\02 - Memorial Descritivo - Com Aumento - {nome_completo}.docx")

            # Alterando formulário para aumento de carga
            if projeto.aumento_carga == True or (projeto.ligacao_nova == True and instalacao.classificacao != "Rural"):
                workbook = load_workbook('C://Users//Henrique Castro//Documents//entrada-dados//workbooks//06 - Formulario para Ligacao Nova ou Aumento de Carga.xlsx')
                sheet = workbook.get_sheet_by_name('Formulário Conexão Nova Urbana')

                sheet["D25"] = cliente.nome
                sheet["D26"] = cliente.nome_pais
                sheet["C27"] = cliente.rg
                sheet["E27"] = cliente.cpf
                sheet["H27"] = cliente.nascimento
                sheet["E28"] = 37998262232
                sheet["H28"] = 3798262232
                sheet["C29"] = "engenharia@apollofotovoltaica.com.br"
                sheet["D30"] = "Não"
                sheet["F31"] = "Não"

                sheet["C34"] = instalacao.logradouro
                sheet["F34"] = instalacao.numero_predial
                sheet["H34"] = instalacao.complemento
                sheet["C35"] = instalacao.bairro
                sheet["F35"] = instalacao.cep
                sheet["C36"] = instalacao.cidade
                sheet["F36"] = "MG"
                sheet["G38"] = "Sim"

                if projeto.ligacao_nova == True:
                    sheet["F41"] = "11.0"
                    sheet["F42"] = "3 - Endereço da Unidade Consumidora"
                
                if projeto.carga == 5:
                    sheet["B74"] = 1
                    sheet["E64"] = 20
                    sheet["E74"] = 1

                elif projeto.carga == 6.5:
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E64"] = 20
                    sheet["E74"] = 1
                    sheet["E80"] = 1

                elif projeto.carga == 10:
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E64"] = 20
                    sheet["E74"] = 1
                    sheet["E76"] = 1
                    sheet["E77"] = 16
                    sheet["E80"] = 1

                elif projeto.carga == 12:
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E65"] = 20
                    sheet["E71"] = 1
                    sheet["E74"] = 1
                    sheet["E77"] = 20
                    sheet["E80"] = 1

                elif projeto.carga == 15 or projeto.carga == 15.1:
                    sheet["E65"] = 20
                    sheet["E71"] = 1
                    sheet["E74"] = 1
                    sheet["E77"] = 19
                    sheet["E78"] = 1
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 16.8:
                    sheet["B82"] = 1
                    sheet["E65"] = 20
                    sheet["E71"] = 1
                    sheet["E74"] = 1
                    sheet["E76"] = 1
                    sheet["E77"] = 22
                    sheet["E78"] = 1
                    sheet["E79"] = 1
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 20:
                    sheet["B82"] = 1
                    sheet["E65"] = 20
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 34
                    sheet["E78"] = 1
                    sheet["E79"] = 2
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 24:
                    sheet["B82"] = 1
                    sheet["E65"] = 35
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 39
                    sheet["E78"] = 1
                    sheet["E79"] = 2
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 30:
                    sheet["B82"] = 1
                    sheet["E65"] = 35
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 39
                    sheet["E78"] = 2
                    sheet["E79"] = 2
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 36:
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E65"] = 35
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 47
                    sheet["E78"] = 2
                    sheet["E79"] = 2
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 50:
                    sheet["B58"] = 3
                    sheet["B74"] = 1
                    sheet["E65"] = 50
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 52
                    sheet["E78"] = 3
                    sheet["E79"] = 4
                    sheet["E80"] = 3
                    sheet["E82"] = 1

                elif projeto.carga == 23:
                    sheet["E65"] = 35
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 39
                    sheet["E78"] = 1
                    sheet["E79"] = 2
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 27:
                    sheet["B74"] = 1
                    sheet["E63"] = 1
                    sheet["E65"] = 35
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 38
                    sheet["E78"] = 1
                    sheet["E79"] = 1
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 38:
                    sheet["B61"] = 1
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E63"] = 1
                    sheet["E65"] = 45
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 2
                    sheet["E76"] = 1
                    sheet["E77"] = 38
                    sheet["E78"] = 2
                    sheet["E79"] = 1
                    sheet["E80"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 47:
                    sheet["B61"] = 1
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E63"] = 1
                    sheet["E65"] = 45
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 3
                    sheet["E76"] = 1
                    sheet["E77"] = 54
                    sheet["E78"] = 2
                    sheet["E79"] = 2
                    sheet["E80"] = 2
                    sheet["E81"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 57:
                    sheet["B61"] = 2
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E63"] = 1
                    sheet["E65"] = 55
                    sheet["E67"] = 30
                    sheet["E71"] = 1
                    sheet["E74"] = 4
                    sheet["E76"] = 4
                    sheet["E77"] = 53
                    sheet["E78"] = 2
                    sheet["E79"] = 3
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 1

                elif projeto.carga == 66:
                    sheet["B61"] = 2
                    sheet["B74"] = 1
                    sheet["B82"] = 1
                    sheet["E63"] = 2
                    sheet["E65"] = 55
                    sheet["E67"] = 30
                    sheet["E69"] = 1
                    sheet["E71"] = 1
                    sheet["E74"] = 4
                    sheet["E76"] = 4
                    sheet["E77"] = 59
                    sheet["E78"] = 3
                    sheet["E79"] = 3
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 1

                elif projeto.carga == 75:
                    sheet["B61"] = 3
                    sheet["B82"] = 1
                    sheet["E63"] = 2
                    sheet["E65"] = 55
                    sheet["E67"] = 30
                    sheet["E69"] = 1
                    sheet["E71"] = 1
                    sheet["E74"] = 4
                    sheet["E76"] = 4
                    sheet["E77"] = 68
                    sheet["E78"] = 4
                    sheet["E79"] = 3
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 2
                
                if projeto.carga >= 50:
                    sheet["D83"] = "Sim"
                    sheet["G83"] = projeto.carga

                else:
                    sheet["D83"] = "Não"

                if projeto.aumento_carga == True:
                    if projeto.novo_n_fases == "Tripolar" or projeto.novo_disjuntor >= 100:
                        sheet["F151"] = "Sim"
                        sheet["B155"] = 1
                        sheet["C155"] = "Motor 2 CV"
                        sheet["D155"] = 1
                        sheet["E155"] = novo_fases + " 220 V"
                        sheet["F155"] = 2
                        sheet["G155"] = "CV"
                        sheet["H155"] = 2.1
                        sheet["D164"] = "Não"

                    else:
                        sheet["F151"] = "Não"

                elif projeto.ligacao_nova == True:
                    if projeto.n_fases == "Tripolar" or projeto.disjuntor >= 100:
                        sheet["F151"] = "Sim"
                        sheet["B155"] = 1
                        sheet["C155"] = "Motor 2 CV"
                        sheet["D155"] = 1
                        sheet["E155"] = fases + " 220 V"
                        sheet["F155"] = 2
                        sheet["G155"] = "CV"
                        sheet["H155"] = 2.1
                        sheet["D164"] = "Não"

                    else:
                        sheet["F151"] = "Não"
                
                n_polos: str = ""
                novo_n_polos: str = ""

                if fases == "Monofásico":
                    n_polos = "1"

                elif fases == "Bifásico":
                    n_polos = "2"
                
                elif fases == "Trifásico":
                    n_polos = "3"

                if novo_fases == "Monofásico":
                    novo_n_polos = "1"

                elif novo_fases == "Bifásico":
                    novo_n_polos = "2"

                elif novo_fases == "Trifásico":
                        novo_n_polos = "3"

                if projeto.aumento_carga == True:
                    sheet["D181"] = "Alteração de Carga"
                    sheet["C186"] = "Instalação de Nº:*"
                    sheet["D186"] = instalacao.numero_instalacao

                    sheet["F186"] = n_polos + "x" + str(projeto.disjuntor) + "A"
                    sheet["H186"] = novo_n_polos + "x" + str(projeto.novo_disjuntor) + "A"
                    sheet["C187"] = instalacao.classificacao
                    sheet["D188"] = "Sim"
                    sheet["F188"] = projeto.carga
                    sheet["H188"] = projeto.carga
                    sheet["D190"] = "Não"

                elif projeto.ligacao_nova == True:
                    sheet["D181"] = "Ligação Nova"
                    sheet["C183"] = projeto.n_fases + " " + str(projeto.disjuntor) + "A"
                    sheet["F183"] = instalacao.numero_predial
                    sheet["H183"] = instalacao.complemento
                    sheet["C184"] = instalacao.classificacao
                    sheet["F184"] = instalacao.classificacao

                workbook.save(f'{diretorio}//06 - Formulario para Ligacao Nova ou Aumento de Carga - {nome_completo}.xlsx')
                workbook.close()

            # Alterando formulário para ligação nova rural
            if projeto.ligacao_nova == True and instalacao.classificacao == "Rural":
                workbook = load_workbook('workbooks//06 - Formulario para Ligacao Nova Rural.xlsx')
                sheet = workbook.get_sheet_by_name('Análise Para Ligação Nova Rural')

                sheet["D23"] = cliente.nome
                sheet["D24"] = cliente.nome_pais
                sheet["C25"] = cliente.rg
                sheet["E25"] = cliente.cpf
                sheet["H25"] = cliente.nascimento
                sheet["E26"] = 37998262232
                sheet["H26"] = 3798262232
                sheet["C27"] = "engenharia@apollofotovoltaica.com.br"
                sheet["D28"] = "Não"
                sheet["F29"] = "Não"

                sheet["D32"] = "11.0"
                sheet["E33"] = "Não"
                
                sheet["C35"] = instalacao.logradouro
                sheet["F35"] = instalacao.numero_predial
                sheet["H35"] = instalacao.complemento
                sheet["C36"] = instalacao.bairro
                sheet["F36"] = instalacao.cep
                sheet["C37"] = instalacao.cidade
                sheet["F37"] = "MG"
                
                latitude = float(str.split(instalacao.coordenadas_decimais, ",")[0])
                longitude = float(str.split(instalacao.coordenadas_decimais, ",")[1])

                sheet["E42"] = latitude
                sheet["H42"] = longitude
                sheet["D45"] = instalacao.bairro
                sheet["C46"] = instalacao.cidade
                sheet["F46"] = instalacao.logradouro
                sheet["D47"] = instalacao.classificacao
                sheet["E48"] = "Não"
                sheet["H48"] = "Não"
                sheet["G53"] = "Sim"

                sheet["B56"] = "Primeiro ponto em propriedade rural."

                sheet["C61"] = fases + " " + str(projeto.disjuntor) + "A"
                
                if projeto.carga == 5:
                    sheet["B75"] = 1
                    sheet["E67"] = 20
                    sheet["E76"] = 1

                elif projeto.carga == 6.5:
                    sheet["B76"] = 1
                    sheet["B79"] = 1
                    sheet["E67"] = 20
                    sheet["E76"] = 1
                    sheet["E82"] = 1

                elif projeto.carga == 10:
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E67"] = 20
                    sheet["E76"] = 1
                    sheet["E78"] = 1
                    sheet["E79"] = 16
                    sheet["E82"] = 1

                elif projeto.carga == 12:
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E68"] = 20
                    sheet["E73"] = 1
                    sheet["E76"] = 1
                    sheet["E79"] = 20
                    sheet["E82"] = 1

                elif projeto.carga == 15 or projeto.carga == 15.1:
                    sheet["E68"] = 20
                    sheet["E73"] = 1
                    sheet["E76"] = 1
                    sheet["E79"] = 19
                    sheet["E80"] = 1
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 16.8:
                    sheet["B79"] = 1
                    sheet["E68"] = 20
                    sheet["E73"] = 1
                    sheet["E76"] = 1
                    sheet["E78"] = 1
                    sheet["E79"] = 22
                    sheet["E80"] = 1
                    sheet["E81"] = 1
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 20:
                    sheet["B79"] = 1
                    sheet["E68"] = 20
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 34
                    sheet["E80"] = 1
                    sheet["E81"] = 2
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 24:
                    sheet["B79"] = 1
                    sheet["E68"] = 35
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 39
                    sheet["E80"] = 1
                    sheet["E81"] = 2
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 30:
                    sheet["B79"] = 1
                    sheet["E68"] = 35
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 39
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 36:
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E68"] = 35
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 47
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 50:
                    sheet["B65"] = 3
                    sheet["B75"] = 1
                    sheet["E68"] = 50
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 52
                    sheet["E80"] = 3
                    sheet["E81"] = 4
                    sheet["E82"] = 3
                    sheet["E84"] = 1

                elif projeto.carga == 23:
                    sheet["E68"] = 35
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 39
                    sheet["E80"] = 1
                    sheet["E81"] = 2
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 27:
                    sheet["B75"] = 1
                    sheet["E66"] = 1
                    sheet["E68"] = 35
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 38
                    sheet["E80"] = 1
                    sheet["E81"] = 1
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 38:
                    sheet["B68"] = 1
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E66"] = 1
                    sheet["E68"] = 45
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 2
                    sheet["E78"] = 1
                    sheet["E79"] = 38
                    sheet["E80"] = 2
                    sheet["E81"] = 1
                    sheet["E82"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 47:
                    sheet["B68"] = 1
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E66"] = 1
                    sheet["E68"] = 45
                    sheet["E70"] = 30
                    sheet["E73"] = 1
                    sheet["E76"] = 3
                    sheet["E78"] = 1
                    sheet["E79"] = 54
                    sheet["E80"] = 2
                    sheet["E81"] = 2
                    sheet["E82"] = 2
                    sheet["E83"] = 1
                    sheet["E84"] = 1

                elif projeto.carga == 57:
                    sheet["B68"] = 2
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E66"] = 2
                    sheet["E68"] = 55
                    sheet["E70"] = 30
                    sheet["E71"] = 1
                    sheet["E73"] = 1
                    sheet["E76"] = 4
                    sheet["E78"] = 4
                    sheet["E79"] = 59
                    sheet["E80"] = 3
                    sheet["E81"] = 3
                    sheet["E82"] = 2
                    sheet["E83"] = 2
                    sheet["E84"] = 1

                elif projeto.carga == 66:
                    sheet["B68"] = 2
                    sheet["B75"] = 1
                    sheet["B79"] = 1
                    sheet["E66"] = 2
                    sheet["E68"] = 55
                    sheet["E70"] = 30
                    sheet["E71"] = 1
                    sheet["E73"] = 1
                    sheet["E76"] = 4
                    sheet["E78"] = 4
                    sheet["E79"] = 59
                    sheet["E80"] = 3
                    sheet["E81"] = 3
                    sheet["E82"] = 2
                    sheet["E83"] = 2
                    sheet["E84"] = 1

                elif projeto.carga == 75:
                    sheet["B68"] = 3
                    sheet["B79"] = 1
                    sheet["E66"] = 2
                    sheet["E68"] = 55
                    sheet["E70"] = 30
                    sheet["E71"] = 1
                    sheet["E73"] = 1
                    sheet["E76"] = 4
                    sheet["E78"] = 4
                    sheet["E79"] = 68
                    sheet["E80"] = 4
                    sheet["E81"] = 3
                    sheet["E82"] = 2
                    sheet["E83"] = 2
                    sheet["E84"] = 2
                
                if projeto.n_fases == "Tripolar" or projeto.disjuntor >= 100:
                    sheet["E87"] = "Sim"
                    sheet["C91"] = fases
                    sheet["D91"] = "220.0"
                    sheet["E91"] = 2
                    sheet["F91"] = "CV"
                    sheet["G91"] = 1
                    sheet["H91"] = 2.1
                    sheet["D100"] = "Não"
                
                sheet["B105"] = "2 - Não tenho interesse em nenhum dos itens oferecidos. Estou ciente que não haverá nenhum tipo de ressarcimento."

                workbook.save(f'{diretorio}//06 - Formulario para Ligacao Nova Rural - {nome_completo}.xlsx')
                workbook.close()

            return f"Formulário do cliente {nome_completo} impresso com sucesso"

        # Caso a ID informada não esteja presente no banco de dados
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f'Projeto com a ID {id} não existe')
        
    # Caso o usuário atual esteja validado
    else:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail=f'Por favor, faça login antes de imprimir um projeto')